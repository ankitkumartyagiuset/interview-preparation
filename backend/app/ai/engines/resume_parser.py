import os
from PyPDF2 import PdfReader
import docx
from typing import Dict, Any, Optional
from app.ai.gateway import ai_gateway
from app.ai.prompts.resume_prompts import RESUME_EXTRACTION_SYSTEM_PROMPT, build_resume_extraction_prompt

class ResumeParserEngine:
    """Extracts raw text from resume documents and generates structured candidate profiles."""

    @staticmethod
    def extract_text_from_file(file_path: str) -> str:
        """Extract plain text from PDF, DOCX, DOC, or TXT file."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        extracted_text = ""

        if ext == ".pdf":
            try:
                reader = PdfReader(file_path)
                pages = [page.extract_text() or "" for page in reader.pages]
                extracted_text = "\n".join(pages)
            except Exception as e:
                raise ValueError(f"Failed to parse PDF document: {e}")
        elif ext in [".docx", ".doc"]:
            try:
                doc = docx.Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text]
                # Also check tables in docx
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            paragraphs.append(" | ".join(row_text))
                extracted_text = "\n".join(paragraphs)
            except Exception as e:
                # If binary doc format fails, try basic fallback
                raise ValueError(f"Failed to parse Word document: {e}")
        elif ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                extracted_text = f.read()
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        extracted_text = extracted_text.strip()
        if not extracted_text:
            raise ValueError("Document appears to be empty or contains no extractable text.")

        return extracted_text

    @staticmethod
    async def extract_structured_profile(raw_text: str, provider_name: Optional[str] = None) -> Dict[str, Any]:
        """Invoke AI Gateway to extract validated structured profile from raw text."""
        prompt = build_resume_extraction_prompt(raw_text)
        structured_data = await ai_gateway.generate_json(
            prompt=prompt,
            system_prompt=RESUME_EXTRACTION_SYSTEM_PROMPT,
            provider_name=provider_name
        )
        return structured_data

resume_parser_engine = ResumeParserEngine()
